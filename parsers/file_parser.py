"""
File parser — extracts text and structured academic data from uploaded files.

Supported formats:
  PDF   → pdfplumber
  DOCX  → python-docx
  XLSX  → openpyxl / pandas
  CSV   → pandas
  TXT   → plain read
  Images → metadata only (no OCR)
"""

from __future__ import annotations

import io
import re
import csv
from dataclasses import dataclass, field
from typing import Optional


# Output types 

@dataclass
class ParsedGrade:
    course_name:  Optional[str]  = None
    course_code:  Optional[str]  = None
    grade_letter: Optional[str]  = None
    score:        Optional[float] = None
    max_score:    Optional[float] = None
    percentage:   Optional[float] = None
    semester:     Optional[str]  = None
    source_row:   Optional[int]  = None


@dataclass
class TextSnippet:
    snippet_type: str   # 'heading' | 'feedback' | 'grade_line' | 'comment'
    content:      str
    page_number:  Optional[int] = None


@dataclass
class ParseResult:
    raw_text:  str = ""
    grades:    list[ParsedGrade]   = field(default_factory=list)
    snippets:  list[TextSnippet]   = field(default_factory=list)
    error:     Optional[str]       = None


# Regex helpers 

# e.g.  CS101, MATH 202, ENG-003
COURSE_CODE_RE = re.compile(r'\b([A-Z]{2,6}[\s\-]?\d{3,4}[A-Z]?)\b')

# Grade letters: A+, A, B-, C+, etc.
GRADE_LETTER_RE = re.compile(r'\b([A-F][+-]?)\b')

# Percentage: 87%, 87.5 %, 87/100
PERCENTAGE_RE = re.compile(r'(\d+\.?\d*)\s*(?:%|/\s*100)\b')

# Score out of max: 42/50  or  42 / 50
SCORE_FRAC_RE = re.compile(r'(\d+\.?\d*)\s*/\s*(\d+\.?\d*)')

# GPA: 3.7/4.0 or 3.70 / 4
GPA_RE = re.compile(r'\b(\d\.\d{1,2})\s*/\s*4\.?0?\b')

# Semester hints: "Semester 1", "Spring 2024", "2023-24 S2"
SEMESTER_RE = re.compile(
    r'(?:semester\s*\d|spring|autumn|fall|summer|winter)\s*\d{0,4}',
    re.IGNORECASE,
)


def _normalise_percentage(score: float, max_score: float) -> float:
    if max_score and max_score > 0:
        return round(score / max_score * 100, 2)
    return score


def _extract_grade_from_line(line: str, row_idx: int = 0) -> Optional[ParsedGrade]:
    """Try to parse a single text line into a ParsedGrade."""
    line = line.strip()
    if len(line) < 3:
        return None

    grade = ParsedGrade(source_row=row_idx)

    # Course code
    cc = COURSE_CODE_RE.search(line)
    if cc:
        grade.course_code = cc.group(1).strip()

    # Percentage first (more specific than bare numbers)
    pct = PERCENTAGE_RE.search(line)
    if pct:
        grade.percentage = float(pct.group(1))
        grade.score      = grade.percentage
        grade.max_score  = 100.0

    # Score fraction: 42/50
    if not grade.percentage:
        frac = SCORE_FRAC_RE.search(line)
        if frac:
            s, m = float(frac.group(1)), float(frac.group(2))
            if m <= 200:   # sanity: avoid matching years like 2023/2024
                grade.score      = s
                grade.max_score  = m
                grade.percentage = _normalise_percentage(s, m)

    # Grade letter
    gl = GRADE_LETTER_RE.search(line)
    if gl:
        grade.grade_letter = gl.group(1)

    # At least one piece of useful data?
    if not any([grade.percentage, grade.score, grade.grade_letter, grade.course_code]):
        return None

    # Heuristic: first "word-like" token before the code/grade is the course name
    name_part = re.split(r'[:\-|]', line)[0].strip()
    name_part  = COURSE_CODE_RE.sub('', name_part).strip()
    name_part  = re.sub(r'\s+', ' ', name_part)
    if 3 <= len(name_part) <= 120:
        grade.course_name = name_part

    # Semester hint
    sem = SEMESTER_RE.search(line)
    if sem:
        grade.semester = sem.group(0).strip()

    return grade


# Per-format extractors

def _parse_txt(content: bytes) -> ParseResult:
    try:
        text = content.decode('utf-8', errors='replace')
    except Exception as e:
        return ParseResult(error=str(e))

    result   = ParseResult(raw_text=text)
    snippets = []
    grades   = []

    for i, line in enumerate(text.splitlines()):
        stripped = line.strip()
        if not stripped:
            continue

        # Heading heuristic: short, ALL CAPS or ends with ':'
        if len(stripped) < 80 and (stripped.isupper() or stripped.endswith(':')):
            snippets.append(TextSnippet('heading', stripped))
            continue

        g = _extract_grade_from_line(stripped, i)
        if g:
            grades.append(g)
        elif len(stripped) > 20:
            snippets.append(TextSnippet('comment', stripped))

    result.grades   = grades
    result.snippets = snippets[:200]   # cap to avoid storing megabytes
    return result


def _parse_pdf(content: bytes) -> ParseResult:
    try:
        import pdfplumber
    except ImportError:
        return ParseResult(error="pdfplumber not installed")

    result   = ParseResult()
    all_text = []
    grades   = []
    snippets = []

    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                all_text.append(page_text)

                # Extract tables if present
                for table in page.extract_tables():
                    if not table:
                        continue
                    headers = [str(c).lower() for c in (table[0] or [])]
                    course_idx = _find_col(headers, ['course', 'module', 'subject', 'unit'])
                    grade_idx  = _find_col(headers, ['grade', 'mark', 'score', 'result', 'percentage'])
                    code_idx   = _find_col(headers, ['code', 'course code', 'module code'])
                    sem_idx    = _find_col(headers, ['semester', 'term', 'year'])

                    for row_i, row in enumerate(table[1:], start=1):
                        if not row:
                            continue
                        g = ParsedGrade(source_row=row_i)
                        if course_idx is not None and course_idx < len(row):
                            g.course_name = str(row[course_idx] or '').strip()
                        if code_idx is not None and code_idx < len(row):
                            g.course_code = str(row[code_idx] or '').strip()
                        if grade_idx is not None and grade_idx < len(row):
                            raw = str(row[grade_idx] or '').strip()
                            g   = _fill_grade_value(g, raw)
                        if sem_idx is not None and sem_idx < len(row):
                            g.semester = str(row[sem_idx] or '').strip()
                        if g.course_name or g.grade_letter or g.percentage:
                            grades.append(g)

                # Line-by-line extraction as fallback
                for i, line in enumerate(page_text.splitlines()):
                    stripped = line.strip()
                    if not stripped:
                        continue
                    if len(stripped) < 80 and stripped.isupper():
                        snippets.append(TextSnippet('heading', stripped, page_num))
                    elif len(stripped) > 30:
                        g = _extract_grade_from_line(stripped, i)
                        if g:
                            grades.append(g)
                        else:
                            snippets.append(TextSnippet('comment', stripped, page_num))

    except Exception as e:
        result.error = str(e)

    result.raw_text = "\n\n".join(all_text)
    result.grades   = _dedup_grades(grades)
    result.snippets = snippets[:300]
    return result


def _parse_docx(content: bytes) -> ParseResult:
    try:
        from docx import Document
    except ImportError:
        return ParseResult(error="python-docx not installed")

    result   = ParseResult()
    all_text = []
    grades   = []
    snippets = []

    try:
        doc = Document(io.BytesIO(content))

        # Paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            all_text.append(text)
            style = para.style.name.lower() if para.style else ''
            if 'heading' in style or (len(text) < 80 and text.isupper()):
                snippets.append(TextSnippet('heading', text))
            else:
                g = _extract_grade_from_line(text, i)
                if g:
                    grades.append(g)
                elif len(text) > 20:
                    snippets.append(TextSnippet('comment', text))

        # Tables
        for table in doc.tables:
            headers = [c.text.lower().strip() for c in table.rows[0].cells] if table.rows else []
            course_idx = _find_col(headers, ['course', 'module', 'subject', 'unit'])
            grade_idx  = _find_col(headers, ['grade', 'mark', 'score', 'result', 'percentage'])
            code_idx   = _find_col(headers, ['code', 'course code'])
            sem_idx    = _find_col(headers, ['semester', 'term'])

            for row_i, row in enumerate(table.rows[1:], start=1):
                cells = [c.text.strip() for c in row.cells]
                g = ParsedGrade(source_row=row_i)
                if course_idx is not None and course_idx < len(cells):
                    g.course_name = cells[course_idx]
                if code_idx is not None and code_idx < len(cells):
                    g.course_code = cells[code_idx]
                if grade_idx is not None and grade_idx < len(cells):
                    g = _fill_grade_value(g, cells[grade_idx])
                if sem_idx is not None and sem_idx < len(cells):
                    g.semester = cells[sem_idx]
                if g.course_name or g.grade_letter or g.percentage:
                    grades.append(g)

    except Exception as e:
        result.error = str(e)

    result.raw_text = "\n".join(all_text)
    result.grades   = _dedup_grades(grades)
    result.snippets = snippets[:300]
    return result


def _parse_csv(content: bytes) -> ParseResult:
    try:
        import pandas as pd
    except ImportError:
        return ParseResult(error="pandas not installed")

    result = ParseResult()
    try:
        df = pd.read_csv(io.BytesIO(content), dtype=str).fillna('')
        result.raw_text = df.to_string()
        result.grades   = _extract_grades_from_df(df)
    except Exception as e:
        result.error = str(e)
    return result


def _parse_xlsx(content: bytes) -> ParseResult:
    try:
        import pandas as pd
    except ImportError:
        return ParseResult(error="pandas not installed")

    result = ParseResult()
    try:
        xls = pd.ExcelFile(io.BytesIO(content))
        all_text_parts = []
        all_grades: list[ParsedGrade] = []
        for sheet in xls.sheet_names:
            df = xls.parse(sheet, dtype=str).fillna('')
            all_text_parts.append(f"--- Sheet: {sheet} ---\n{df.to_string()}")
            all_grades.extend(_extract_grades_from_df(df))
        result.raw_text = "\n\n".join(all_text_parts)
        result.grades   = all_grades
    except Exception as e:
        result.error = str(e)
    return result


# Shared helpers 

COURSE_KEYWORDS   = ['course', 'module', 'subject', 'class', 'unit', 'paper']
GRADE_KEYWORDS    = ['grade', 'mark', 'score', 'result', 'percentage', 'gpa', 'pct', 'pts', 'points']
CODE_KEYWORDS     = ['code', 'course code', 'module code', 'subject code']
SEMESTER_KEYWORDS = ['semester', 'term', 'year', 'period']


def _find_col(headers: list[str], keywords: list[str]) -> Optional[int]:
    for i, h in enumerate(headers):
        if any(kw in h for kw in keywords):
            return i
    return None


def _fill_grade_value(g: ParsedGrade, raw: str) -> ParsedGrade:
    raw = raw.strip()
    if not raw:
        return g

    # Grade letter
    gl = GRADE_LETTER_RE.match(raw)
    if gl and len(raw) <= 3:
        g.grade_letter = gl.group(1)
        return g

    # Percentage
    pct = PERCENTAGE_RE.search(raw)
    if pct:
        g.percentage = float(pct.group(1))
        g.score, g.max_score = g.percentage, 100.0
        return g

    # Fraction
    frac = SCORE_FRAC_RE.search(raw)
    if frac:
        s, m = float(frac.group(1)), float(frac.group(2))
        if m <= 200:
            g.score, g.max_score = s, m
            g.percentage = _normalise_percentage(s, m)
        return g

    # Bare number
    try:
        val = float(raw)
        if val <= 4.0:          # treat as GPA
            g.percentage = round(val / 4.0 * 100, 2)
        elif val <= 100:
            g.percentage = val
            g.score, g.max_score = val, 100.0
        else:
            g.score = val
    except ValueError:
        pass

    return g


def _extract_grades_from_df(df) -> list[ParsedGrade]:
    import pandas as pd
    grades = []
    headers_lower = [str(c).lower().strip() for c in df.columns]

    course_idx = _find_col(headers_lower, COURSE_KEYWORDS)
    grade_idx  = _find_col(headers_lower, GRADE_KEYWORDS)
    code_idx   = _find_col(headers_lower, CODE_KEYWORDS)
    sem_idx    = _find_col(headers_lower, SEMESTER_KEYWORDS)

    cols = list(df.columns)

    for row_i, row in df.iterrows():
        g = ParsedGrade(source_row=int(row_i) + 2)  # +2 accounts for header row

        if course_idx is not None:
            g.course_name = str(row.iloc[course_idx]).strip()
        if code_idx is not None:
            g.course_code = str(row.iloc[code_idx]).strip()
        if grade_idx is not None:
            g = _fill_grade_value(g, str(row.iloc[grade_idx]))
        if sem_idx is not None:
            g.semester = str(row.iloc[sem_idx]).strip()

        # Fallback: if no column detected, scan all cells for grade-like values
        if grade_idx is None:
            for val in row.values:
                g = _fill_grade_value(g, str(val))
                if g.grade_letter or g.percentage:
                    break

        if g.course_name or g.grade_letter or g.percentage:
            grades.append(g)

    return grades


def _dedup_grades(grades: list[ParsedGrade]) -> list[ParsedGrade]:
    """Remove exact duplicates (same course + same score)."""
    seen = set()
    out  = []
    for g in grades:
        key = (g.course_name, g.course_code, g.grade_letter, g.percentage)
        if key not in seen:
            seen.add(key)
            out.append(g)
    return out


# Public entry point

def parse_file(file_type: str, content: bytes) -> ParseResult:
    """
    Dispatch to the correct parser based on file extension.
    Returns a ParseResult with raw_text, grades, and snippets.
    """
    ft = file_type.lower()
    if ft == 'pdf':
        return _parse_pdf(content)
    if ft in ('doc', 'docx'):
        return _parse_docx(content)
    if ft == 'csv':
        return _parse_csv(content)
    if ft in ('xlsx', 'xls'):
        return _parse_xlsx(content)
    if ft == 'txt':
        return _parse_txt(content)
    # Images and unknown formats — no text extraction
    return ParseResult(raw_text="", error=None)
